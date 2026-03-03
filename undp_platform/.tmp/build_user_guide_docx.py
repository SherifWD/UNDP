from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root = Path('/Applications/XAMPP/xamppfiles/htdocs/UNDP/undp_platform')
out = root / 'docs' / 'UNDP_Platform_User_Guide.docx'
shots = root / 'docs' / 'screenshots'

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Arial'

bookmark_id = 1

def add_bookmark(paragraph, name):
    global bookmark_id
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    bookmark_id += 1


def add_internal_link(paragraph, text, anchor):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1D4ED8')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_title(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Title']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run.bold = True
    return p


def add_heading(level, text, bookmark=None):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def add_image(file_name, caption):
    path = shots / file_name
    if not path.exists():
        add_para(f'[Screenshot missing: {caption}]')
        return
    doc.add_picture(str(path), width=Inches(6.6))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x5F, 0x6B, 0x85)

# Cover + TOC
add_title('UNDP Platform User Guide and User Manual')
add_para('Platform: UNDP Monitoring / Validation Platform (Web Dashboard, APIs, and Mobile Integration Surface)')
add_para('Generated from the current codebase on 27 February 2026.')
add_para('This document is structured with numbered sections and a clickable contents page. Use the links below to jump to the section you need.')

doc.add_paragraph('Guide Contents', style='Heading 1')
contents = [
    ('1. System Overview', 'section_1'),
    ('2. Roles and Permissions', 'section_2'),
    ('3. Seeded Test Accounts', 'section_3'),
    ('4. Workflow and Lifecycle Stages', 'section_4'),
    ('5. Common Layout and Navigation', 'section_5'),
    ('6. Page-by-Page Guide', 'section_6'),
    ('6.1 Login', 'section_6_1'),
    ('6.2 OTP Verification', 'section_6_2'),
    ('6.3 Dashboard', 'section_6_3'),
    ('6.4 Projects', 'section_6_4'),
    ('6.5 Project Submissions', 'section_6_5'),
    ('6.6 Validation Worklist', 'section_6_6'),
    ('6.7 Submission Detail', 'section_6_7'),
    ('6.8 Reports & Map', 'section_6_8'),
    ('6.9 Partner Dashboard', 'section_6_9'),
    ('6.10 User Management', 'section_6_10'),
    ('6.11 Audit Log', 'section_6_11'),
    ('6.12 Settings', 'section_6_12'),
    ('6.13 Municipal Overview', 'section_6_13'),
    ('6.14 Access Denied', 'section_6_14'),
    ('7. Role-Based Scenario Playbooks', 'section_7'),
    ('8. Error Handling and Support', 'section_8'),
    ('9. Mobile API Companion', 'section_9'),
    ('10. Deliverables in This Repository', 'section_10'),
]
for label, anchor in contents:
    p = doc.add_paragraph(style='List Bullet')
    add_internal_link(p, label, anchor)

doc.add_page_break()

# Sections
add_heading(1, '1. System Overview', 'section_1')
add_para('The platform is a role-aware monitoring and validation system used to manage projects, receive field evidence, process validation decisions, run reports, maintain audit trails, and administer configuration.')

add_heading(1, '2. Roles and Permissions', 'section_2')
add_table(
    ['Role', 'Main Responsibility', 'Typical Access'],
    [
        ['Reporter', 'Create and track own submissions.', 'Dashboard, Projects, media upload, own data'],
        ['Municipal Focal Point', 'Validate municipality submissions.', 'Dashboard, Validation, Reports & Map'],
        ['UNDP Admin', 'Manage the full system.', 'All admin pages and exports'],
        ['Partner / Donor Viewer', 'Read-only approved/aggregated review.', 'Partner dashboard, approved aggregated data'],
        ['Auditor', 'Inspect immutable audit trails.', 'Audit Log and system reporting'],
    ]
)

add_heading(1, '3. Seeded Test Accounts', 'section_3')
add_table(
    ['Role', 'Name', 'Email', 'Phone'],
    [
        ['UNDP Admin', 'UNDP Admin', 'admin@undp.local', '+218910000001'],
        ['Auditor', 'Audit Officer', 'auditor@undp.local', '+218910000002'],
        ['Municipal Focal Point', 'Municipal Focal Point - Tripoli', 'focal.tripoli@undp.local', '+218910000003'],
        ['Partner / Donor Viewer', 'Partner Donor Viewer', 'partner@undp.local', '+218910000004'],
        ['Municipal Focal Point', 'Municipal Focal Point - Benghazi', 'focal.benghazi@undp.local', '+218910000005'],
        ['Municipal Focal Point', 'Municipal Focal Point - Misrata', 'focal.misrata@undp.local', '+218910000006'],
        ['Reporter', 'Community Reporter - Tripoli', 'reporter.tripoli@undp.local', '+218910000101'],
        ['Reporter', 'Community Reporter - Benghazi', 'reporter.benghazi@undp.local', '+218910000102'],
        ['Reporter', 'Community Reporter - Misrata', 'reporter.misrata@undp.local', '+218910000103'],
    ]
)
add_para('In local development, OTPs are typically written to storage/logs/laravel.log when the log OTP sender is active.')

add_heading(1, '4. Workflow and Lifecycle Stages', 'section_4')
add_heading(2, '4.1 Submission Stages')
add_numbered(['Draft', 'Queued', 'Submitted', 'Under Review', 'Approved', 'Rework Requested', 'Rejected'])
add_heading(2, '4.2 Operational Stages')
add_numbered(['Authentication', 'Project Selection', 'Submission Creation', 'Validation', 'Reporting / Audit / Export', 'Administrative Maintenance'])

add_heading(1, '5. Common Layout and Navigation', 'section_5')
add_bullets([
    'The left sidebar is permission-aware and only shows routes allowed for the logged-in role.',
    'The top header shows the current page label, language switcher, and notification button.',
    'The sidebar footer card shows the logged-in user and the logout action.',
    'Most tables use pagination with Prev, numbered page buttons, and Next.',
    'Unauthorized navigation attempts route to the Access Denied screen.'
])
add_image('03-dashboard.png', 'Live screenshot: authenticated application shell with sidebar, header, and dashboard context')

doc.add_page_break()
add_heading(1, '6. Page-by-Page Guide', 'section_6')

add_heading(2, '6.1 Login Page (/login)', 'section_6_1')
add_para('Purpose: start the OTP authentication flow.')
add_bullets(['Fields: Country Code, Phone Number', 'Primary button: Continue', 'Common outcomes: successful OTP request, invalid phone validation, resend cooldown'])
add_image('01-login.png', 'Login page with country code and phone number entry')

add_heading(2, '6.2 OTP Verification Page (/otp)', 'section_6_2')
add_para('Purpose: complete login using a 6-digit code.')
add_bullets(['Controls: OTP input boxes, Verify, Resend, Change Number', 'Common outcomes: successful login, invalid code, expired code, disabled account'])
add_image('02-otp.png', 'OTP verification page with six input boxes and verification controls')

add_heading(2, '6.3 Dashboard (/)', 'section_6_3')
add_para('Purpose: KPI overview with municipality/project map, search, filtering, and project drill-down.')
add_bullets([
    'Buttons and controls: Add New Project, dropdown filters, Municipality, Search projects, Filter, Apply, Reset.',
    'Project cards in the right rail are clickable and open the project detail pane.',
    'When a project is selected, the map board shifts and the details pane opens on the right.'
])
add_image('03-dashboard.png', 'Dashboard page showing KPI panels and the municipality/projects board')

add_heading(2, '6.4 Projects Page (/projects)', 'section_6_4')
add_para('Purpose: browse, create, edit, and organize projects.')
add_bullets(['Controls: Add New Project, Search, Status filter, Municipality tabs, View Submission, Edit, pagination'])
add_image('04-projects.png', 'Projects page showing the project registry table')

add_heading(2, '6.5 Project Submissions Page (/projects/:id/submissions)', 'section_6_5')
add_para('Purpose: project-specific submission overview and export surface.')
add_bullets(['Controls: Back, Status filter, Search submissions, Export CSV, Export PDF, View Details'])
add_image('05-project-submissions.png', 'Project submissions page with project KPIs and submission rows')

add_heading(2, '6.6 Validation Worklist (/validation)', 'section_6_6')
add_para('Purpose: pending review queue for validators and municipal focal points.')
add_bullets(['Controls: Newest, Oldest, By Project, search, filters, Reset, Apply, row navigation'])
add_image('06-validation.png', 'Validation worklist page used to process pending submissions')

add_heading(2, '6.7 Submission Detail (/submissions/:id)', 'section_6_7')
add_para('Purpose: inspect one submission and execute status transitions.')
add_bullets(['Controls: Approve, Request Rework, Reject, media preview buttons, action modal confirm/cancel'])
add_para('This page is implemented but was not auto-captured in the current screenshot batch.')

add_heading(2, '6.8 Reports & Map (/reports)', 'section_6_8')
add_para('Purpose: analytics, map inspection, chart drill-down, and exports.')
add_bullets(['Controls: date range, municipality, project, status, Apply, Reset Filters, Export CSV, Export PDF, Download export, Full Screen Map'])
add_image('10-reports-map.png', 'Reports & Map page showing reporting filters, charts, and map area')

add_heading(2, '6.9 Partner Dashboard (/partner-dashboard)', 'section_6_9')
add_para('Purpose: read-only, partner-focused summary page for approved or aggregated reporting.')
add_para('The route is implemented and permission-gated but was not auto-captured in this batch.')

add_heading(2, '6.10 User Management (/users)', 'section_6_10')
add_para('Purpose: central user administration.')
add_bullets(['Controls: Add New User, Set Permissions, Accept/Enable, Disable, pagination, modal save/cancel actions'])
add_image('07-users.png', 'User Management page with summary cards and user action rows')

add_heading(2, '6.11 Audit Log (/audit-logs)', 'section_6_11')
add_para('Purpose: immutable system activity review.')
add_bullets(['Controls: action filter, user/date filters, Today, Last 7 Days, Reset, Apply, row detail'])
add_image('08-audit-log.png', 'Audit Log page with summary counters, filters, and log table')

add_heading(2, '6.12 Settings (/settings)', 'section_6_12')
add_para('Purpose: persisted configuration management across General, Users & Roles, Reporting & Workflow, and Security.')
add_bullets(['Controls: tab buttons, Reset, Save Changes, list editors, policy toggles, numeric fields'])
add_image('09-settings.png', 'Settings page showing tabs and persisted system configuration controls')

add_heading(2, '6.13 Municipal Overview (/municipal-overview)', 'section_6_13')
add_para('This page is implemented and route-protected but currently hidden from the sidebar. It remains available as a direct deep-link.')

add_heading(2, '6.14 Access Denied (/access-denied)', 'section_6_14')
add_para('Users are redirected here when a route exists but the current role lacks the required permission.')

add_heading(1, '7. Role-Based Scenario Playbooks', 'section_7')
add_heading(2, '7.1 Reporter')
add_numbered(['Login with OTP.', 'Open Dashboard and locate a project.', 'Create a submission from mobile/API flow.', 'Track status changes and respond to rework.'])
add_heading(2, '7.2 Municipal Focal Point')
add_numbered(['Open Validation Worklist.', 'Filter pending items.', 'Open a submission.', 'Approve, reject, or request rework.', 'Review dashboard and reports.'])
add_heading(2, '7.3 UNDP Admin')
add_numbered(['Manage projects, municipalities, and users.', 'Use Settings for workflow and security.', 'Use Audit Log for traceability.', 'Run exports and reporting.'])
add_heading(2, '7.4 Partner / Donor Viewer')
add_numbered(['Open Partner Dashboard.', 'Filter and review read-only metrics.', 'Export summaries when permitted.'])
add_heading(2, '7.5 Auditor')
add_numbered(['Open Audit Log.', 'Filter actor, action, or date.', 'Inspect before/after payloads and metadata.'])

add_heading(1, '8. Error Handling and Support', 'section_8')
add_table(
    ['Condition', 'User Impact', 'Recommended Action'],
    [
        ['OTP resend too soon', '429 cooldown', 'Wait for cooldown and retry.'],
        ['Invalid / expired OTP', '422 validation error', 'Request a new code.'],
        ['Disabled account', '403 login blocked', 'Re-enable from User Management.'],
        ['Permission denied', '403 or Access Denied', 'Check role and permissions.'],
        ['422 validation on forms', 'Save blocked', 'Check required fields and referenced IDs.'],
        ['Export not ready', 'Task still queued/processing', 'Poll export task until ready.'],
    ]
)

add_heading(1, '9. Mobile API Companion', 'section_9')
add_bullets([
    'Use docs/postman/UNDP_Mobile_API.postman_collection.json.',
    'Use docs/postman/UNDP_Mobile_API_Local.postman_environment.json.',
    'The collection covers OTP auth, profile, projects, municipalities, submissions, media, dashboard, audit, settings, and exports.',
    'In local development, OTP can be read from storage/logs/laravel.log when the log provider is active.'
])

add_heading(1, '10. Deliverables in This Repository', 'section_10')
add_bullets([
    'docs/UNDP_Platform_User_Guide.docx',
    'docs/UNDP_Platform_User_Guide.html',
    'public/UNDP_Platform_User_Guide.html',
    'docs/postman/UNDP_Mobile_API.postman_collection.json',
    'docs/postman/UNDP_Mobile_API_Local.postman_environment.json',
    'docs/screenshots/*.png and public/screenshots/*.png'
])

doc.save(out)
print(out)
