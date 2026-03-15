from __future__ import annotations
from datetime import date
from pathlib import Path
from html import escape
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/Applications/XAMPP/xamppfiles/htdocs/UNDP/undp_platform')
DOCS = ROOT / 'docs'
PUBLIC = ROOT / 'public'
SHOTS = DOCS / 'screenshots'
PUBLIC_SHOTS = PUBLIC / 'screenshots'
DOCX_OUT = DOCS / 'UNDP_Platform_User_Guide.docx'
HTML_DOCS_OUT = DOCS / 'UNDP_Platform_User_Guide.html'
HTML_PUBLIC_OUT = PUBLIC / 'UNDP_Platform_User_Guide.html'
GENERATED_ON = date.today().strftime('%d %B %Y')

PUBLIC_SHOTS.mkdir(parents=True, exist_ok=True)
for img in SHOTS.glob('*.png'):
    copy2(img, PUBLIC_SHOTS / img.name)

roles_table = {
    'headers': ['Role', 'Main Responsibility', 'Primary Screens'],
    'rows': [
        ['Reporter', 'Creates and tracks field submissions.', 'Dashboard, Projects, Project Submissions (scoped), own submission detail'],
        ['Municipal Focal Point', 'Validates municipality evidence, reviews scoped submissions, and monitors municipality KPIs.', 'Dashboard, Validation Worklist, Submission Detail, Project Submissions, Municipal Overview, KPI & Geo Reports'],
        ['UNDP Admin', 'Full platform administration, donor funding review, and system oversight.', 'Dashboard, Projects, Project Detail, Users, Audit Log, Settings, KPI & Geo Reports, Municipal Overview'],
        ['Auditor', 'Reviews immutable activity history.', 'Audit Log, exports, reporting'],
        ['Partner / Donor Viewer', 'Reads approved aggregated information and submits project funding requests.', 'Partner dashboard, Projects, Project Detail, Funding Request modal, approved reporting views'],
    ],
}

crud_table = {
    'headers': ['Module', 'Create', 'Read', 'Update', 'Delete / Archive'],
    'rows': [
        ['Municipalities', 'Web UI supports create via Projects page modal.', 'Visible in filters, tabs, dashboard selectors, and overview screens.', 'API supports update; dedicated web edit screen is not currently exposed.', 'Delete is not implemented in the current web UI/API surface.'],
        ['Projects', 'Web UI supports create.', 'Web UI supports list and detail modal.', 'Web UI supports edit.', 'Web UI supports delete from the project detail modal for authorized admins.'],
        ['Submissions', 'Creation is handled by the mobile app / API flow, not by the current admin web UI.', 'Web UI supports list, project list, worklist, and detail view.', 'Status transitions (approve / rework / reject) are supported; direct content editing is not exposed in the web UI.', 'Delete is not implemented in the current web UI/API surface.'],
        ['Funding Requests', 'Donor users can create requests from the Projects table and project detail modal.', 'Donors can read their project request history; admins can review all requests in Projects and KPI & Geo Reports.', 'Admins can approve or decline with a mandatory review reason.', 'Delete is not exposed; keep the audited decision history instead.'],
        ['Users', 'Web UI supports create.', 'Web UI supports list and audit lookups.', 'Web UI supports role / municipality / profile updates and status toggling.', 'Hard delete is not implemented; use disable / enable instead.'],
        ['Audit Logs', 'System-generated only.', 'Web UI supports list and entry detail.', 'Immutable by design.', 'Not allowed.'],
        ['Settings', 'System settings record is created automatically on first access.', 'Web UI supports full read.', 'Web UI supports save / reset by tab.', 'Not allowed.'],
    ],
}

entries = [
    {
        'anchor': 'section-1',
        'title': '1. Guide Scope and Navigation',
        'paragraphs': [
            'This guide covers the currently implemented web dashboard, donor funding flow, focal-point validation flow, reporting screens, user administration, settings, and the API surface used by the mobile application. It is written against the live codebase and the current seeded dataset.',
            'The cover page works as a clickable table of contents. In the HTML guide, each link jumps directly to the relevant section. In the Word guide, the same numbering and internal bookmarks are preserved.',
        ],
        'callouts': [('notice', 'All walkthroughs below reflect the current implementation. Admin, Donor, and Municipal Focal Point journeys are documented with refreshed screenshots from the live dashboard instead of older generic placeholders.')],
    },
    {
        'anchor': 'section-2',
        'title': '2. Role Matrix',
        'paragraphs': ['Use this matrix to understand which role normally performs each workflow. Permission checks are enforced in both the UI and the API.'],
        'table': roles_table,
    },
    {
        'anchor': 'section-3',
        'title': '3. Workflow and Status Lifecycle',
        'paragraphs': ['Every evidence item and donor funding request moves through a defined lifecycle. Review actions are audited and cannot be silently altered.'],
        'subsections': [
            {
                'anchor': 'section-3-1',
                'title': '3.1 Submission Statuses',
                'numbered': ['Draft', 'Queued', 'Submitted', 'Under Review', 'Approved', 'Rework Requested', 'Rejected'],
            },
            {
                'anchor': 'section-3-2',
                'title': '3.2 Operational Stages',
                'numbered': ['Authenticate', 'Navigate to the target project or work queue', 'Open the target record', 'Perform review or administration action', 'Verify result in audit / reporting surfaces'],
            },
            {
                'anchor': 'section-3-3',
                'title': '3.3 Funding Request Statuses',
                'numbered': ['Pending Review', 'Approved', 'Declined'],
            },
        ],
    },
    {
        'anchor': 'section-4',
        'title': '4. CRUD Availability Matrix',
        'paragraphs': ['The platform still restricts hard delete for most core records, but projects and funding requests now have clearer operational paths in the web dashboard. Use this matrix as the authoritative reference when training users.'],
        'table': crud_table,
    },
    {
        'anchor': 'section-5',
        'title': '5. Authentication Module',
        'paragraphs': ['Authentication is phone-number based and uses OTP verification. The same token is reused by the web app and can be used in Postman/mobile during development.'],
        'subsections': [
            {
                'anchor': 'section-5-1',
                'title': '5.1 Login Page (/login)',
                'purpose': 'Start the OTP request flow.',
                'buttons': ['Country code selector', 'Phone number field', 'Continue'],
                'steps': [
                    'Open the login page.',
                    'Select the country code if needed. The default local setup uses +218.',
                    'Enter the phone number and click Continue.',
                    'If validation passes, the app navigates to the OTP page.',
                    'If validation fails, the field shows an inline error and the request is blocked.',
                ],
                'images': [('01-login.png', 'Login page with country code selector and phone number field')],
            },
            {
                'anchor': 'section-5-2',
                'title': '5.2 OTP Verification (/otp)',
                'purpose': 'Complete the login and create the authenticated session.',
                'buttons': ['OTP input fields', 'Verify', 'Resend', 'Change Number'],
                'steps': [
                    'Enter the 6-digit OTP from the configured SMS / log provider.',
                    'Click Verify.',
                    'On success, the app stores the token and user object locally, then redirects to the dashboard.',
                    'Use Resend if the code expires or is not received.',
                    'Use Change Number to return to the login page and restart the flow.',
                ],
                'images': [('02-otp.png', 'OTP verification screen with six input slots and verification controls')],
            },
            {
                'anchor': 'section-5-3',
                'title': '5.3 Logout and Session Control',
                'purpose': 'Terminate the session safely from the sidebar footer card.',
                'steps': [
                    'Locate the logged-in user card at the bottom of the sidebar.',
                    'Click the logout action on the right side of the card.',
                    'The stored token is removed and the app returns to the login page.',
                ],
                'images': [('03-dashboard.png', 'Authenticated layout showing the sidebar user card used for logout')],
            },
        ],
    },
    {
        'anchor': 'section-6',
        'title': '6. Admin Dashboard Module',
        'paragraphs': ['The Home dashboard combines KPIs with a municipality and project workspace. This is the main landing page for internal dashboard users such as UNDP Admin and other non-partner staff.', 'Partner / Donor users are redirected to the read-only partner dashboard after login, while Municipal Focal Point users rely on Validation Worklist and Municipal Overview for their main operational journey.'],
        'subsections': [
            {
                'anchor': 'section-6-1',
                'title': '6.1 KPI Strip and Global Overview',
                'purpose': 'Review headline metrics before drilling into projects.',
                'buttons': ['All Projects dropdown', 'All Sources dropdown', 'All Time dropdown', 'Add New Project'],
                'steps': [
                    'Open the dashboard after login.',
                    'Review the Reports donut, Funding Progress block, and Beneficiaries Overview block.',
                    'Use the top-level dropdowns to scope visible metrics where available.',
                    'Use Add New Project to jump into project creation if your role has permission.',
                ],
                'images': [('03-dashboard.png', 'Dashboard top section with KPI cards and global actions')],
            },
            {
                'anchor': 'section-6-2',
                'title': '6.2 Municipality and Project Workspace',
                'purpose': 'Browse projects by municipality and map them spatially.',
                'buttons': ['Municipality dropdown', 'Search projects', 'Filter', 'Project list cards'],
                'steps': [
                    'Use the Municipality dropdown to scope the map and project list. Leaving it on All Municipalities keeps the full scoped dataset visible.',
                    'Use Search projects to filter the visible project rail by name or reference.',
                    'Review the map markers and the live project count in the right rail.',
                    'Select a project card to drill into the detail pane and inspect project-specific counts before leaving the dashboard.',
                ],
                'images': [
                    ('03-dashboard.png', 'Dashboard workspace in its default wide state'),
                    ('11-dashboard-selected-project.png', 'Dashboard workspace after selecting a project and opening the detail pane'),
                ],
            },
            {
                'anchor': 'section-6-3',
                'title': '6.3 Filter Panel and Drill-Down',
                'purpose': 'Refine the project rail using priority, area, and status.',
                'buttons': ['Filter', 'Apply', 'Reset', 'Go to Submission', 'Close details'],
                'steps': [
                    'Click Filter next to the search field.',
                    'Choose priority, area, and status as needed.',
                    'Click Apply to update the project rail and map focus.',
                    'Click Reset to clear only the dashboard-side project filters.',
                    'After opening a project, use Go to Submission to move directly into the project-specific submission page for deeper review.',
                ],
                'images': [
                    ('12-dashboard-filter-panel.png', 'Dashboard filter panel open'),
                    ('11-dashboard-selected-project.png', 'Dashboard project detail pane with direct submission navigation'),
                ],
            },
        ],
    },
    {
        'anchor': 'section-7',
        'title': '7. Municipality Management',
        'paragraphs': ['Municipalities are currently managed from the Projects module and then reused throughout filters, tabs, and dashboards.'],
        'subsections': [
            {
                'anchor': 'section-7-1',
                'title': '7.1 Create a Municipality',
                'purpose': 'Add a municipality record used by projects and scoped dashboards.',
                'steps': [
                    'Open Projects.',
                    'Click the Create Municipality button in the page header.',
                    'Enter English name, Arabic name, and the municipality code.',
                    'Click Save Municipality to persist the new record.',
                    'The municipality becomes available in project forms, tabs, and dashboard selectors after reload.',
                ],
                'images': [('16-municipality-create-modal.png', 'Create Municipality modal in the Projects module')],
            },
            {
                'anchor': 'section-7-2',
                'title': '7.2 Read and Use Municipalities',
                'purpose': 'Use municipality records to segment data.',
                'steps': [
                    'In Projects, municipality tabs split the project list by municipality.',
                    'In Dashboard and Reports & Map, municipality selectors scope the visible markers and project lists.',
                    'In Settings and User Management, municipality assignment determines what some roles can review.',
                ],
                'images': [
                    ('04-projects.png', 'Projects page showing municipality tabs'),
                    ('27-municipal-overview.png', 'Municipal overview page used to inspect scoped summary data'),
                ],
                'callouts': [('warning', 'Municipality update exists at the API layer, but the current web UI does not expose a dedicated municipality edit screen. Hard delete is not implemented.')],
            },
        ],
    },
    {
        'anchor': 'section-8',
        'title': '8. Project Management',
        'paragraphs': ['Projects are the core container for reporting and donor funding activity. The web UI now supports create, list, detail, edit, delete, and donor funding-request entry points depending on role.'],
        'subsections': [
            {
                'anchor': 'section-8-1',
                'title': '8.1 Browse the Project Registry',
                'purpose': 'Read project rows and key metrics.',
                'buttons': ['Search', 'Status filter', 'Municipality tabs', 'Request to Fund', 'View Submissions', 'Edit'],
                'steps': [
                    'Open Projects from the sidebar.',
                    'Use search and status filter to narrow the rows.',
                    'Use municipality tabs to limit the table to one municipality.',
                    'Review approved submissions, pending submissions, progress, current execution status, and funding request summary directly in the table.',
                    'If you are logged in as a donor, use the Request to Fund button directly in the row to start a funding request without opening the project first.',
                ],
                'images': [('04-projects.png', 'Projects registry table')],
            },
            {
                'anchor': 'section-8-2',
                'title': '8.2 Create a Project',
                'purpose': 'Create a new project record.',
                'steps': [
                    'Click Add New Project in the Projects header.',
                    'Select the municipality and enter English and Arabic names.',
                    'Enter a description, choose status, and set latitude/longitude.',
                    'Click Save Project.',
                    'After save, verify that the new row appears in the registry and on the dashboard map.',
                ],
                'images': [('13-project-create-modal.png', 'Create Project modal with project fields')],
            },
            {
                'anchor': 'section-8-3',
                'title': '8.3 View Project Details',
                'purpose': 'Open the project detail modal for the full operational summary.',
                'steps': [
                    'Click any project row in the Projects table.',
                    'Review the project status, municipality, last update, location, description, progress, funding information, and submission stats.',
                    'Use Go to Submissions to jump to the project-specific submission page.',
                    'If you are a donor, scroll to the Project Funding Requests section to inspect request history and use the in-detail request button.',
                ],
                'images': [('14-project-detail-modal.png', 'Project detail modal opened from the project registry')],
            },
            {
                'anchor': 'section-8-4',
                'title': '8.4 Edit a Project',
                'purpose': 'Update project metadata.',
                'steps': [
                    'Click Edit in a project row, or open the detail modal and click Edit there.',
                    'Modify any allowed field: municipality, names, description, status, latitude, or longitude.',
                    'Click Update Project to save the change.',
                    'Re-open the detail modal to confirm the saved values.',
                ],
                'images': [('15-project-edit-modal.png', 'Edit Project modal')],
            },
            {
                'anchor': 'section-8-5',
                'title': '8.5 Donor Request to Fund from the Projects Table',
                'purpose': 'Submit a funding request from the fastest entry point in the donor journey.',
                'steps': [
                    'Log in as a Partner / Donor Viewer and open Projects.',
                    'Locate the target project row and review the funding summary column.',
                    'Click Request to Fund in that row.',
                    'Enter the requested amount and optional donor note, then submit the modal.',
                    'The new request appears in project-level history and becomes available to admin review.',
                ],
                'images': [('30-donor-projects-request-button.png', 'Donor project table showing the row-level Request to Fund action')],
            },
            {
                'anchor': 'section-8-6',
                'title': '8.6 Donor Request to Fund from Project Detail',
                'purpose': 'Submit or review funding requests inside the project detail modal.',
                'steps': [
                    'Open the target project row to load the detail modal.',
                    'Scroll to the Project Funding Requests section.',
                    'Review total requested amount, pending count, approved / declined count, and prior donor request reasons.',
                    'Click Request to Fund This Project to open the funding modal from inside the detail view.',
                ],
                'images': [
                    ('31-donor-project-detail-funding.png', 'Donor project detail showing the Project Funding Requests section'),
                    ('32-donor-funding-request-modal.png', 'Donor funding request modal opened from project detail'),
                ],
            },
            {
                'anchor': 'section-8-7',
                'title': '8.7 Delete a Project',
                'purpose': 'Remove a project and its dependent assignments when you have admin permission.',
                'steps': [
                    'Open the project detail modal as a UNDP Admin.',
                    'Click Delete in the modal header.',
                    'Read the confirmation warning carefully because related submissions and assignments are removed with the project.',
                    'Confirm the browser dialog to complete the deletion.',
                    'Verify the project row disappears from the registry after reload.',
                ],
            },
        ],
    },
    {
        'anchor': 'section-9',
        'title': '9. Submission Management',
        'paragraphs': ['Submission records are reviewed in the web UI. Submission creation itself belongs to the mobile or direct API flow.'],
        'subsections': [
            {
                'anchor': 'section-9-1',
                'title': '9.1 Project Submission List',
                'purpose': 'Inspect submission counts and row-level records for one project.',
                'buttons': ['Back', 'Status filter', 'Search submissions', 'Export CSV', 'Export PDF', 'View Details'],
                'steps': [
                    'From Projects, click View Submission for the target project.',
                    'Review the KPI strip for total submissions, pending actions, and evidence coverage.',
                    'Use the table to inspect reporter name, report type, region, attachments, and status.',
                    'Use export buttons to generate the current view in CSV or PDF.',
                ],
                'images': [('05-project-submissions.png', 'Project-level submissions page with KPI cards and row table')],
            },
            {
                'anchor': 'section-9-2',
                'title': '9.2 Open Submission Detail',
                'purpose': 'Read the complete submission, media assets, and timeline.',
                'steps': [
                    'Click View Details on any submission row.',
                    'Review the header details: status, project, municipality, reporter, description, and validation comment.',
                    'Inspect media evidence and click any asset to load a preview.',
                    'Review the immutable timeline at the bottom of the page.',
                ],
                'images': [('20-submission-detail.png', 'Submission detail page with evidence panel and timeline')],
            },
            {
                'anchor': 'section-9-3',
                'title': '9.3 Approve, Rework, and Reject',
                'purpose': 'Process a submission through the validation workflow.',
                'steps': [
                    'Open a submission detail page.',
                    'Choose Approve, Request Rework, or Reject.',
                    'For rework and rejection, provide a comment. Use a reason template where available.',
                    'Click the confirm button in the action modal.',
                    'The status updates and the timeline is extended with the new event.',
                ],
                'images': [('21-submission-action-modal.png', 'Submission action modal used for rework / rejection / approval confirmation')],
            },
            {
                'anchor': 'section-9-4',
                'title': '9.4 Submission Creation Path',
                'purpose': 'Clarify where new submissions are created.',
                'steps': [
                    'The current admin web UI does not expose a Create Submission form.',
                    'New submissions are expected to be created by the mobile application or through direct API integration.',
                    'Use the Postman collection in this repository as the reference for the Flutter/mobile integration team.',
                ],
                'callouts': [('notice', 'This separation is intentional: the web app is currently focused on review, oversight, administration, and reporting, while field submission entry belongs to mobile/API flows.')],
            },
        ],
    },
    {
        'anchor': 'section-10',
        'title': '10. Validation Worklist',
        'paragraphs': ['The validation worklist is the main Municipal Focal Point queue. It is scoped to the assigned municipality and keeps pending review work visible without exposing unrelated submissions.'],
        'buttons': ['Newest', 'Oldest', 'By Project', 'Search', 'Status filter', 'Project filter', 'Date filters', 'Apply', 'Reset', 'Review'],
        'steps': [
            'Open Validation Worklist from the sidebar.',
            'Confirm the municipality scope banner at the top of the page.',
            'Use sorting, project, status, and date filters to isolate the queue you need.',
            'Click Review on the target row to open Submission Detail.',
            'Approve, request rework, or reject the submission, then return to the queue for the next item.',
        ],
        'images': [('06-validation.png', 'Municipal Focal Point validation worklist with scoped pending submissions')],
    },
    {
        'anchor': 'section-11',
        'title': '11. KPI & Geo Reports',
        'paragraphs': ['KPI & Geo Reports is the main analytics surface for dashboard roles that can access reporting. It combines richer KPI cards, municipality and project breakdowns, trend charts, backlog aging, funding charts, export actions, and the interactive map.'],
        'subsections': [
            {
                'anchor': 'section-11-1',
                'title': '11.1 Analytics, Filters, and Map',
                'purpose': 'Review system-wide or scoped analytics before exporting or drilling into details.',
                'buttons': ['Date From', 'Date To', 'Municipality filter', 'Project filter', 'Status filter', 'Report type', 'Apply', 'Reset Filters', 'Export CSV', 'Export PDF', 'Download', 'Full Screen Map'],
                'steps': [
                    'Open KPI & Geo Reports from the sidebar.',
                    'Set the date range and optional municipality, project, status, and report-type filters.',
                    'Click Apply to refresh KPI cards, status analytics, backlog aging, municipality and project breakdowns, trend charts, funding overview, and the map.',
                    'Use chart segments or the interactive map to explain system performance during review meetings or exports.',
                    'Use Export CSV or Export PDF to generate the current filtered output when your role has export permission.',
                ],
                'images': [('10-reports-map.png', 'KPI & Geo Reports analytics page with charts, filters, and map')],
            },
            {
                'anchor': 'section-11-2',
                'title': '11.2 Admin Funding Request Review',
                'purpose': 'Approve or decline donor funding requests with a mandatory review reason.',
                'buttons': ['Pending Review filter', 'Approved filter', 'Declined filter', 'Refresh Requests', 'Approve', 'Decline'],
                'steps': [
                    'Scroll to Funding Requests Review (Admin) in the reports page.',
                    'Keep the filter on Pending Review to work the active donor queue first.',
                    'Read the project, municipality, donor, amount, request reason, and requested timestamp.',
                    'Enter the review reason in the text area.',
                    'Click Approve or Decline. The request status, review timestamp, and funding charts refresh after the decision.',
                ],
                'images': [('33-admin-reports-funding-review.png', 'Admin funding request review section inside KPI & Geo Reports')],
            },
        ],
    },
    {
        'anchor': 'section-12',
        'title': '12. User Management',
        'paragraphs': ['User Management covers onboarding, role assignment, municipality assignment, and account status control.'],
        'subsections': [
            {
                'anchor': 'section-12-1',
                'title': '12.1 User Overview and Read',
                'purpose': 'Review user distribution and the current user registry.',
                'steps': [
                    'Open User Management from the sidebar.',
                    'Review the reporter donut, municipality distribution, and user classification bar.',
                    'Use the table to inspect name, email, phone, role, and current status actions.',
                ],
                'images': [('07-users.png', 'User Management overview and user table')],
            },
            {
                'anchor': 'section-12-2',
                'title': '12.2 Add a New User',
                'purpose': 'Create a new user account and assign the base role.',
                'steps': [
                    'Click Add New User.',
                    'Enter full name, email, country code, and phone number.',
                    'Choose the user role, organization, and assigned municipality if applicable.',
                    'Click Add to create the account.',
                    'The new user appears in the table after the list reloads.',
                ],
                'images': [('17-user-add-modal.png', 'Add New User modal')],
            },
            {
                'anchor': 'section-12-3',
                'title': '12.3 Update Role, Municipality, and Permission Profile',
                'purpose': 'Edit an existing user without recreating the account.',
                'steps': [
                    'Click Set Permissions on the target row.',
                    'Update the username, email, phone, role, or municipality assignment.',
                    'Use the permission preview grid to confirm the effective role-based permission set.',
                    'Click Save Changes.',
                ],
                'images': [('18-user-permissions-modal.png', 'Set Permissions modal for editing an existing user')],
            },
            {
                'anchor': 'section-12-4',
                'title': '12.4 Disable or Re-enable a User',
                'purpose': 'Control access without deleting the account.',
                'steps': [
                    'Use the Disable button on the target row to suspend access.',
                    'Provide a reason if desired and click Confirm.',
                    'To restore access later, use the Enable button on the same row.',
                    'This action is logged to the audit trail.',
                ],
                'images': [('19-user-status-modal.png', 'User status modal for disable / enable actions')],
                'callouts': [('warning', 'Hard delete is not exposed in the current implementation. Disabling the account is the supported control path.')],
            },
        ],
    },
    {
        'anchor': 'section-13',
        'title': '13. Audit Log',
        'paragraphs': ['Audit logs are immutable records of key actions across the platform. They support filtering, review, and forensic traceability.'],
        'subsections': [
            {
                'anchor': 'section-13-1',
                'title': '13.1 Audit Table and Filters',
                'purpose': 'Review high-level activity and find target events.',
                'steps': [
                    'Open Audit Log from the sidebar.',
                    'Use action, user ID, and date filters to narrow the dataset.',
                    'Use Today or Last 7 Days for quick filters.',
                    'Click Apply to refresh the table, or Reset to clear the filter set.',
                ],
                'images': [('08-audit-log.png', 'Audit Log table and summary counters')],
            },
            {
                'anchor': 'section-13-2',
                'title': '13.2 Open Audit Entry Detail',
                'purpose': 'Inspect the before, after, and metadata payload for one event.',
                'steps': [
                    'Click any audit row or its reference link.',
                    'Review the Before payload, After payload, and Metadata block.',
                    'Use this view when confirming who changed what and when.',
                ],
                'images': [('22-audit-log-detail-modal.png', 'Audit entry detail modal')],
                'callouts': [('notice', 'Audit records are view-only and intentionally immutable. There is no create, edit, or delete action in this module.')],
            },
        ],
    },
    {
        'anchor': 'section-14',
        'title': '14. Settings Module',
        'paragraphs': ['The Settings page is now backed by the persistent system settings API. It supports save and reset behavior by tab for authorized users.'],
        'subsections': [
            {
                'anchor': 'section-14-1',
                'title': '14.1 General Tab',
                'purpose': 'Maintain organization profile, localization defaults, and system defaults.',
                'steps': [
                    'Open Settings and stay on the General tab.',
                    'Update organization profile fields, localization defaults, and reporting defaults as needed.',
                    'Use the Add/Edit/Delete mini-controls in status and risk option sets where shown.',
                    'Click Save Changes to persist the active tab.',
                ],
                'images': [('23-settings-general.png', 'Settings page - General tab')],
            },
            {
                'anchor': 'section-14-2',
                'title': '14.2 Users & Roles Tab',
                'purpose': 'Review role definitions and the permission matrix.',
                'steps': [
                    'Click the Users & Roles tab.',
                    'Review the role definition blocks and the core permissions table.',
                    'Confirm the permission matrix before changing user roles in User Management.',
                    'Save changes if you adjust any editable control in this tab.',
                ],
                'images': [('24-settings-users-roles.png', 'Settings page - Users & Roles tab')],
            },
            {
                'anchor': 'section-14-3',
                'title': '14.3 Reporting & Workflow Tab',
                'purpose': 'Configure workflow mode, deadlines, escalation, evidence rules, and notification behavior.',
                'steps': [
                    'Click the Reporting & Workflow tab.',
                    'Set workflow mode, default submission status, approval requirements, escalation rules, and evidence requirements.',
                    'Adjust email and in-app notification toggles as required by policy.',
                    'Click Save Changes after completing the workflow configuration.',
                ],
                'images': [('25-settings-reporting-workflow.png', 'Settings page - Reporting & Workflow tab')],
            },
            {
                'anchor': 'section-14-4',
                'title': '14.4 Security Tab',
                'purpose': 'Manage security posture for authentication and audit controls.',
                'steps': [
                    'Click the Security tab.',
                    'Review or update two-factor authentication, SSO, password rules, RBAC toggle, IP restrictions, and audit toggles.',
                    'Use Save Changes to persist the security configuration.',
                ],
                'images': [('26-settings-security.png', 'Settings page - Security tab')],
            },
            {
                'anchor': 'section-14-5',
                'title': '14.5 Save and Reset Behavior',
                'purpose': 'Use settings actions safely.',
                'steps': [
                    'Each tab can be edited independently.',
                    'Save Changes writes the currently active tab payload to the backend settings record.',
                    'Reset restores the visible tab values from the last saved backend state.',
                    'Users without workflow.manage can still view settings, but cannot persist changes.',
                ],
            },
        ],
    },
    {
        'anchor': 'section-15',
        'title': '15. Role-Specific Dashboards and Access Control',
        'paragraphs': ['These screens complete the main Admin, Donor, and Municipal Focal Point dashboard journeys. Include them in training because they explain where each role lands and what they can access.'],
        'subsections': [
            {
                'anchor': 'section-15-1',
                'title': '15.1 Municipal Overview (/municipal-overview)',
                'purpose': 'Inspect scoped municipal analytics and operational summaries.',
                'steps': [
                    'Open Municipal Overview as a Municipal Focal Point or an admin with municipality dashboard access.',
                    'Use the municipality selector, search field, and status legend to scope the view.',
                    'Review KPI cards, the municipal status donut, and the project list with direct links to project details.',
                ],
                'images': [('27-municipal-overview.png', 'Municipal Overview page for municipality-scoped dashboard users')],
            },
            {
                'anchor': 'section-15-2',
                'title': '15.2 Partner / Donor Read-Only Dashboard (/partner-dashboard)',
                'purpose': 'Explain the donor landing page and its approved-data-only analytics.',
                'steps': [
                    'After login, partner-only users are redirected here instead of the internal home dashboard.',
                    'Use date, municipality, project, and search filters to review approved aggregated data.',
                    'Review the My Funding Requests chart, approved submission breakdowns, municipality bars, project bars, and approved trend.',
                    'Export approved data if CSV or PDF access is enabled for that donor role.',
                ],
                'images': [('29-partner-dashboard.png', 'Partner / Donor read-only dashboard with approved-data KPIs and funding request summary')],
            },
            {
                'anchor': 'section-15-3',
                'title': '15.3 Access Denied (/access-denied)',
                'purpose': 'Explain what users see when a route is blocked by RBAC.',
                'steps': [
                    'If a user attempts to open a route without the required permission, the router redirects to Access Denied.',
                    'The screen explains the issue and offers a button to return home.',
                    'This is expected behavior and confirms that RBAC is being enforced correctly.',
                ],
                'images': [('28-access-denied.png', 'Access Denied screen')],
            },
        ],
    },
    {
        'anchor': 'section-16',
        'title': '16. Mobile API Quick Start',
        'paragraphs': [
            'The mobile application should use the Postman collection in docs/postman as the starting contract. The current web UI already proves the OTP auth, project lookup, submission review, media, and settings endpoints.',
            'Recommended first mobile flow: request OTP, verify OTP, call /auth/me, load municipalities, load scoped projects, create submissions via the API, upload media via presign/complete, then poll the user submission list for status changes.',
        ],
        'bullets': [
            'Collection: docs/postman/UNDP Mobile API.postman_collection.json',
            'Environment: docs/postman/UNDP_Mobile_API_Local.postman_environment.json',
            'Base URL (local): http://127.0.0.1:8000/api',
        ],
    },
    {
        'anchor': 'section-17',
        'title': '17. Repository Deliverables',
        'bullets': [
            'HTML guide (docs copy): docs/UNDP_Platform_User_Guide.html',
            'HTML guide (public copy): public/UNDP_Platform_User_Guide.html',
            'Word guide: docs/UNDP_Platform_User_Guide.docx',
            'Screenshots (docs): docs/screenshots/*.png',
            'Screenshots (public): public/screenshots/*.png',
            'Postman collection: docs/postman/UNDP Mobile API.postman_collection.json',
        ],
    },
]

# ---------- HTML ----------
STYLE = """
body { font-family: -apple-system, BlinkMacSystemFont, 'Helvetica Neue', Arial, sans-serif; color:#1f2940; line-height:1.55; margin:32px; }
h1, h2, h3, h4 { color:#14213d; margin-top:24px; }
h1 { font-size: 28px; border-bottom:2px solid #dfe4ef; padding-bottom:10px; }
h2 { font-size: 22px; border-bottom:1px solid #e6ebf4; padding-bottom:6px; }
h3 { font-size: 18px; }
h4 { font-size: 14px; margin-bottom:6px; }
p, li, td, th { font-size: 11px; }
a { color:#1d4ed8; text-decoration:none; }
a:hover { text-decoration:underline; }
ul, ol { margin-top: 6px; }
.small { color:#5f6b85; font-size:10px; }
.cover { page-break-after: always; border:1px solid #dbe3f1; border-radius:14px; padding:28px; background:#fafcff; }
.cover h1 { border-bottom:none; margin-top:0; }
.toc { columns:2; column-gap:32px; margin-top:16px; }
.toc ul { break-inside: avoid; list-style: none; padding-left:0; margin:0 0 14px; }
.toc li { margin:4px 0; }
.toc .level-1 { font-weight:700; margin-top:10px; }
.toc .level-2 { margin-left:14px; font-weight:400; }
.table { width:100%; border-collapse: collapse; margin:12px 0 18px; }
.table th, .table td { border:1px solid #d9dfeb; padding:8px; vertical-align:top; }
.table th { background:#f3f6fb; text-align:left; }
.callout { border:1px solid #d8dfef; background:#f8faff; padding:12px; border-radius:8px; margin:12px 0; }
.notice { border-left:4px solid #1d4ed8; padding:10px 12px; background:#eef4ff; margin:12px 0; }
.warning { border-left:4px solid #c2410c; padding:10px 12px; background:#fff7ed; margin:12px 0; }
.code { font-family: Menlo, Monaco, monospace; font-size:10px; background:#f4f6fa; padding:2px 4px; border-radius:4px; }
.figure { margin:14px 0 24px; }
.guide-shot { width:100%; max-width:980px; border:1px solid #d5ddea; border-radius:12px; display:block; }
figcaption { font-size:10px; color:#5f6b85; margin-top:6px; }
.back-link { display:inline-block; margin-top:8px; font-size:10px; }
.section-break { page-break-before: always; }
.subgrid { margin-left: 8px; }
"""


def render_table(table: dict) -> str:
    out = ['<table class="table">', '<tr>']
    out.extend(f'<th>{escape(h)}</th>' for h in table['headers'])
    out.append('</tr>')
    for row in table['rows']:
        out.append('<tr>')
        out.extend(f'<td>{escape(cell)}</td>' for cell in row)
        out.append('</tr>')
    out.append('</table>')
    return ''.join(out)


def render_images(images: list[tuple[str, str]]) -> str:
    html = []
    for file_name, caption in images:
        html.append(
            f'<figure class="figure"><img class="guide-shot" src="screenshots/{escape(file_name)}" alt="{escape(caption)}"><figcaption>{escape(caption)}</figcaption></figure>'
        )
    return ''.join(html)


def render_entry_html(entry: dict, heading_tag: str = 'h2') -> str:
    parts = [f'<{heading_tag} id="{escape(entry["anchor"])}">{escape(entry["title"])}</{heading_tag}>']
    for paragraph in entry.get('paragraphs', []):
        parts.append(f'<p>{escape(paragraph)}</p>')
    for kind, text in entry.get('callouts', []):
        css = 'notice' if kind == 'notice' else 'warning'
        parts.append(f'<div class="{css}">{escape(text)}</div>')
    if entry.get('purpose'):
        parts.append(f'<p><strong>Purpose:</strong> {escape(entry["purpose"])}</p>')
    if entry.get('buttons'):
        parts.append('<h4>Main buttons and controls</h4><ul>')
        parts.extend(f'<li>{escape(item)}</li>' for item in entry['buttons'])
        parts.append('</ul>')
    if entry.get('bullets'):
        parts.append('<ul>')
        parts.extend(f'<li>{escape(item)}</li>' for item in entry['bullets'])
        parts.append('</ul>')
    if entry.get('numbered'):
        parts.append('<ol>')
        parts.extend(f'<li>{escape(item)}</li>' for item in entry['numbered'])
        parts.append('</ol>')
    if entry.get('steps'):
        parts.append('<h4>Operational steps</h4><ol>')
        parts.extend(f'<li>{escape(item)}</li>' for item in entry['steps'])
        parts.append('</ol>')
    if entry.get('table'):
        parts.append(render_table(entry['table']))
    if entry.get('images'):
        parts.append(render_images(entry['images']))
    if entry.get('subsections'):
        for sub in entry['subsections']:
            parts.append('<div class="subgrid">')
            parts.append(render_entry_html(sub, 'h3'))
            parts.append('</div>')
    parts.append('<a class="back-link" href="#top">Back to Contents</a>')
    return ''.join(parts)


def build_toc(entries: list[dict]) -> tuple[str, list[tuple[str, str, int]]]:
    items: list[tuple[str, str, int]] = []
    for entry in entries:
        items.append((entry['title'], entry['anchor'], 1))
        for sub in entry.get('subsections', []):
            items.append((sub['title'], sub['anchor'], 2))
    html = ['<div class="toc"><ul>']
    for title, anchor, level in items:
        css = 'level-1' if level == 1 else 'level-2'
        html.append(f'<li class="{css}"><a href="#{escape(anchor)}">{escape(title)}</a></li>')
    html.append('</ul></div>')
    return ''.join(html), items


toc_html, toc_items = build_toc(entries)
html_parts = [
    '<!DOCTYPE html><html lang="en"><head><meta charset="utf-8">',
    '<title>UNDP Platform User Guide</title>',
    f'<style>{STYLE}</style>',
    '</head><body>',
    '<div class="cover" id="top">',
    '<h1>UNDP Platform User Guide and User Manual</h1>',
    '<p><strong>Platform:</strong> UNDP Monitoring / Validation Platform (Web Dashboard, APIs, and Mobile Integration Surface)</p>',
    f'<p><strong>Generated from current codebase:</strong> {escape(GENERATED_ON)}</p>',
    '<p>This edition refreshes the Admin, Donor, and Municipal Focal Point journeys with current screenshots, updated donor funding-request walkthroughs, and explicit notes wherever a capability is API-only or intentionally not exposed in the current web UI.</p>',
    '<h2>Guide Contents</h2>',
    toc_html,
    '</div>',
]
for i, entry in enumerate(entries):
    if i > 0 and entry['anchor'] in {'section-5', 'section-8', 'section-12', 'section-14'}:
        html_parts.append('<div class="section-break"></div>')
    html_parts.append(render_entry_html(entry, 'h2'))
html_parts.append('</body></html>')
html = '\n'.join(html_parts)
HTML_DOCS_OUT.write_text(html, encoding='utf-8')
HTML_PUBLIC_OUT.write_text(html, encoding='utf-8')

# ---------- DOCX ----------
doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    styles[style_name].font.name = 'Arial'

bookmark_id = 1


def add_bookmark(paragraph, name: str):
    global bookmark_id
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    bookmark_id += 1


def add_internal_link(paragraph, text: str, anchor: str):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1D4ED8')
    rpr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)
    new_run.append(rpr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(level: int, text: str, bookmark: str | None = None):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_para(text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(items: list[str]):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(items: list[str]):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_table_doc(table: dict):
    t = doc.add_table(rows=1, cols=len(table['headers']))
    t.style = 'Table Grid'
    for i, h in enumerate(table['headers']):
        t.rows[0].cells[i].text = h
    for row in table['rows']:
        cells = t.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = cell


def add_image_doc(file_name: str, caption: str):
    path = SHOTS / file_name
    if not path.exists():
        add_para(f'[Screenshot missing: {caption}]')
        return
    doc.add_picture(str(path), width=Inches(6.6))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x5F, 0x6B, 0x85)


def render_doc_entry(entry: dict, level: int = 1):
    add_heading(level, entry['title'], entry['anchor'])
    for paragraph in entry.get('paragraphs', []):
        add_para(paragraph)
    for kind, text in entry.get('callouts', []):
        prefix = 'Note: ' if kind == 'notice' else 'Warning: '
        add_para(text, prefix)
    if entry.get('purpose'):
        add_para(entry['purpose'], 'Purpose: ')
    if entry.get('buttons'):
        add_para('Main buttons and controls:', None)
        add_bullets(entry['buttons'])
    if entry.get('bullets'):
        add_bullets(entry['bullets'])
    if entry.get('numbered'):
        add_numbered(entry['numbered'])
    if entry.get('steps'):
        add_para('Operational steps:')
        add_numbered(entry['steps'])
    if entry.get('table'):
        add_table_doc(entry['table'])
    if entry.get('images'):
        for file_name, caption in entry['images']:
            add_image_doc(file_name, caption)
    for sub in entry.get('subsections', []):
        render_doc_entry(sub, min(level + 1, 3))
    p = doc.add_paragraph()
    add_internal_link(p, 'Back to Contents', 'top')


title = doc.add_paragraph()
title.style = doc.styles['Title']
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run('UNDP Platform User Guide and User Manual')
rt.bold = True
rt.font.name = 'Arial'
rt.font.size = Pt(22)
subtitle = doc.add_paragraph('Role-based operational guide with refreshed Admin, Donor, and Municipal Focal Point screenshots.')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_toc_heading = add_heading(1, 'Guide Contents', 'top')
for label, anchor, _level in toc_items:
    p = doc.add_paragraph(style='List Bullet')
    add_internal_link(p, label, anchor)
doc.add_page_break()

for idx, entry in enumerate(entries):
    if idx and entry['anchor'] in {'section-5', 'section-8', 'section-12', 'section-14'}:
        doc.add_page_break()
    render_doc_entry(entry, 1)

DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(DOCX_OUT)
print('generated html and docx')
